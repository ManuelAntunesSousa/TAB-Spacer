import os
from docx import Document

# Path to the input file
input_file = '.docx'
output_file = '.docx'

# Check if the input file exists
if not os.path.exists(input_file):
    print(f"Error: The file '{input_file}' was not found.")
else:
    try:
        # Load the Word document
        doc = Document(input_file)

        # Create a new document for the output
        new_doc = Document()

        # Process each paragraph in the document
        for para in doc.paragraphs:
            new_text = '\t'.join(para.text)  # Add a tab between each character
            new_doc.add_paragraph(new_text)

        # Save the new document
        new_doc.save(output_file)

        print(f"Tabbing completed successfully. The output file is '{output_file}'.")
    except Exception as e:
        print(f"An error occurred: {e}")
